from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DARK = RGBColor(31, 31, 31)
LIGHT_GRAY = "E7E6E6"
MID_GRAY = "D9E1F2"
NOTE_GREEN = "E2F0D9"
WHITE = RGBColor(255, 255, 255)


SCENARIO_LABELS = {
    "baseline": "Referência",
    "lights_off": "Luzes desligadas",
    "shielded_70pct": "Blindagem de 70%",
    "amber_proxy": "Espectro âmbar (proxy)",
    "bright_moon_proxy": "Lua mais intensa (proxy)",
    "combined_mitigation": "Mitigação combinada",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def set_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    tr_pr.append(cant)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, before, after in [
        ("Title", 23, 0, 6),
        ("Heading 1", 17, 11, 5),
        ("Heading 2", 13, 9, 4),
        ("Heading 3", 11, 7, 3),
    ]:
        st = styles[name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = DARK
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Caption"].font.name = "Arial"
    styles["Caption"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Caption"].font.size = Pt(9)
    styles["Caption"].font.italic = True
    styles["Caption"].font.color.rgb = DARK
    styles["Caption"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Caption"].paragraph_format.space_before = Pt(2)
    styles["Caption"].paragraph_format.space_after = Pt(7)

    header = section.header.paragraphs[0]
    header.text = "Modelo espacial estocástico de orientação de filhotes"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(100, 100, 100)
    add_page_number(section.footer.paragraphs[0])
    for run in section.footer.paragraphs[0].runs:
        run.font.name = "Arial"
        run.font.size = Pt(8.5)


def add_text(doc: Document, text: str, *, bold_prefix: str | None = None, italic=False, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        r = p.add_run(text)
        r.italic = italic


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)


def add_formula(doc: Document, tokens: Iterable[tuple[str, str | None]]) -> None:
    """Adiciona fórmula editável como texto matemático, com subscritos horizontais."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    for text, mode in tokens:
        r = p.add_run(text)
        r.font.name = "Cambria Math"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
        r.font.size = Pt(11.5)
        if mode == "sub":
            r.font.subscript = True
        elif mode == "sup":
            r.font.superscript = True
        elif mode == "italic":
            r.italic = True


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for j, text in enumerate(headers):
        cell = hdr.cells[j]
        set_cell_shading(cell, MID_GRAY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(text))
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(font_size)
    for row_data in rows:
        row = table.add_row()
        set_cant_split(row)
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 or len(str(text)) > 20 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(text))
            r.font.name = "Arial"
            r.font.size = Pt(font_size)
    if widths:
        for row in table.rows:
            for j, width in enumerate(widths):
                row.cells[j].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_note_box(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, NOTE_GREEN)
    set_cell_margins(cell, top=120, bottom=120, start=140, end=140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(10.5)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    for r2 in p2.runs:
        r2.font.name = "Arial"
        r2.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc: Document, image: Path, caption: str, width: float = 6.2) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run()
    run.add_picture(str(image), width=Inches(width))
    doc.add_paragraph(caption, style="Caption")
    # Mantém apenas a imagem unida à legenda. Não encadeamos a legenda ao
    # elemento seguinte, pois isso pode empurrar várias figuras para outra página.
    set_keep_with_next(p)


def add_reference(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.28)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(10.3)


def fmt_pct(v: float) -> str:
    return f"{100*v:.1f}%".replace(".", ",")


def fmt_num(v: float, nd=2) -> str:
    return f"{v:.{nd}f}".replace(".", ",")


def generate(project: Path, output: Path) -> None:
    scenario = pd.read_csv(project / "outputs_demo" / "experiments" / "scenario_summary.csv")
    calibration = json.loads((project / "outputs_demo" / "calibration.json").read_text(encoding="utf-8"))
    behavior_validation = json.loads((project / "outputs_demo" / "behavior_validation.json").read_text(encoding="utf-8"))
    light_validation = json.loads((project / "outputs_demo" / "light_validation.json").read_text(encoding="utf-8"))
    sensitivity = pd.read_csv(project / "outputs_demo" / "sensitivity" / "sensitivity_spearman.csv")
    figs = project / "outputs_demo" / "experiments" / "figures"

    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph(style="Title")
    title.add_run("Modelo Espacial Estocástico Calibrável para a Orientação de Filhotes de Tartarugas Marinhas sob Poluição Luminosa")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    r = subtitle.add_run("Estudo de caso ambiental com medições direcionais de Delray Beach")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(12.5)

    add_note_box(
        doc,
        "Status científico desta versão",
        "O campo luminoso foi parametrizado com medições publicadas de Delray Beach. O perfil transversal, os ninhos e os obstáculos do exemplo são demonstrativos, e a calibração comportamental utiliza trajetórias sintéticas para verificar o método. Assim, os resultados numéricos apresentados nesta versão não constituem estimativas ecológicas finais da praia.",
    )

    doc.add_paragraph("Resumo", style="Heading 3")
    add_text(doc, "Este trabalho apresenta uma reformulação completa de um autômato celular estocástico destinado à simulação da orientação de filhotes de tartarugas marinhas durante a travessia da praia. O modelo anterior utilizava um campo escalar cuja implementação fazia a suposta poluição luminosa reforçar o movimento em direção ao oceano. A nova formulação substitui esse mecanismo por pistas vetoriais concorrentes, fisicamente escaladas em metros e segundos, associadas ao horizonte marítimo, ao brilho artificial direcional, à silhueta da duna, à declividade e à persistência do deslocamento. O campo luminoso do estudo de caso foi parametrizado a partir de nove locais de Delray Beach, com medições nas direções duna, norte, oceano, sul e zênite. A escolha entre os movimentos da vizinhança de Moore é efetuada por uma função Softmax aplicada à utilidade de cada direção. O fluxo computacional inclui calibração por evolução diferencial, validação ambiental por exclusão de locais, validação comportamental em dados independentes, Monte Carlo por cenário e análise global de sensibilidade. Na demonstração sintética, a condição de referência apresentou chegada ao mar média de 66,3%, enquanto os cenários de luzes desligadas e mitigação combinada produziram 94,6% e 95,0%, respectivamente. Esses valores demonstram o comportamento interno do modelo, mas dependem de geometria assumida e parâmetros sintéticos. A utilização preditiva requer levantamento topográfico, ninhos, radiometria e trajetórias observadas no mesmo setor de praia.")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.add_run("Palavras-chave: ").bold = True
    p.add_run("autômato celular; modelo baseado em agentes; tartarugas marinhas; poluição luminosa; calibração; Monte Carlo; orientação animal.")

    doc.add_paragraph("1. Introdução", style="Heading 1")
    add_text(doc, "A travessia entre o ninho e a linha d’água é uma etapa espacialmente curta, porém biologicamente crítica, da história de vida das tartarugas marinhas. Os filhotes emergem predominantemente à noite e utilizam um conjunto de pistas visuais para selecionar a direção do oceano. Em praias naturais, a maior radiância no horizonte marítimo, a menor silhueta nessa direção e a declividade da praia contribuem para a orientação. Medições espectrais em praias naturais mostraram radiância superior no horizonte marítimo em relação ao terrestre, enquanto estudos de campo associaram o contraste entre as direções terrestre e marítima e variáveis lunares à ocorrência de desorientação (Salmon et al., 1992; Celano et al., 2018; Stanley et al., 2020).")
    add_text(doc, "A iluminação artificial altera a distribuição angular do brilho. Por isso, a variável relevante não é somente uma intensidade escalar, mas a magnitude relativa das pistas, sua direção, visibilidade, espectro e interação com o perfil da duna. Um coeficiente abstrato de zero a três, desacompanhado de unidade ou calibração, não pode ser interpretado diretamente como lux, radiância ou risco biológico.")
    add_text(doc, "O artigo original propôs uma energia total dada pela diferença entre uma atração marítima e uma componente artificial. Entretanto, as definições computacionais adotadas faziam ambos os termos favorecerem o aumento da coordenada transversal, isto é, o avanço em direção ao oceano. A intensidade artificial e a sensibilidade comportamental também apareciam parcialmente confundidas no produto que determinava a Softmax. A presente versão corrige essa inconsistência e reorganiza o trabalho como um modelo calibrável e verificável.")

    doc.add_paragraph("1.1 Pergunta e hipóteses", style="Heading 2")
    add_text(doc, "Pergunta central: como o contraste e a direção das pistas luminosas, combinados à geometria da praia e à variabilidade individual, alteram a orientação, o tempo de exposição e a probabilidade de um filhote alcançar a linha d’água?")
    add_bullet(doc, "H1 — O aumento da componente artificial terrestre reduz a chegada ao mar e aumenta a dispersão angular.")
    add_bullet(doc, "H2 — Fontes de mesma magnitude produzem trajetórias diferentes conforme a direção e a oclusão pela duna.")
    add_bullet(doc, "H3 — A persistência reduz mudanças bruscas de rumo, mas pode prolongar uma direção inicialmente inadequada.")
    add_bullet(doc, "H4 — Blindagem, redução da componente artificial e fortalecimento da pista marítima diminuem o tempo de travessia.")

    doc.add_paragraph("2. Área de estudo e dados de entrada", style="Heading 1")
    doc.add_paragraph("2.1 Levantamento luminoso de Delray Beach", style="Heading 2")
    add_text(doc, "Hirama et al. (2022) publicaram um levantamento realizado em Delray Beach, Flórida, em 17 de agosto de 2020, um dia antes da lua nova. Foram efetuadas oito leituras consecutivas em quatro direções horizontais e no zênite, em nove locais distribuídos ao longo de aproximadamente 800 m, totalizando 360 medições. O conjunto incluído no código reproduz as médias das leituras 5 a 8 apresentadas pelos autores. A unidade do Sky Quality Meter é magnitude por segundo de arco quadrado; valores menores representam maior brilho.")
    add_figure(doc, figs / "light_profile.png", "Figura 1. Brilho direcional usado na calibração ambiental. O eixo vertical foi invertido porque menor magnitude corresponde a maior brilho. Fonte dos dados: Hirama et al. (2022).", 6.25)

    doc.add_paragraph("2.2 Hierarquia de evidências", style="Heading 2")
    add_table(
        doc,
        ["Componente", "Estado nesta versão", "Uso permitido"],
        [
            ["Brilho direcional", "Medido e publicado", "Calibração ambiental relativa"],
            ["Posição longitudinal dos locais", "Derivada das coordenadas", "Interpolação ao longo da costa"],
            ["Largura, altura da duna e declive", "Hipótese demonstrativa", "Teste do programa; substituir"],
            ["Ninhos e obstáculos", "Hipótese demonstrativa", "Teste do programa; substituir"],
            ["Trajetórias comportamentais", "Sintéticas", "Verificação da calibração; não inferência"],
            ["Predação e energia", "Parâmetros de cenário", "Análise de sensibilidade"],
        ],
        [1.65, 1.75, 2.95],
    )
    add_text(doc, "Essa separação impede que valores criados para a demonstração sejam confundidos com observações de campo. Uma versão preditiva deve substituir todos os componentes demonstrativos por dados obtidos no mesmo setor, período e sistema de coordenadas.")

    doc.add_paragraph("3. Descrição do modelo", style="Heading 1")
    add_text(doc, "A descrição segue a lógica do protocolo ODD — Overview, Design concepts and Details — atualizado por Grimm et al. (2020). O modelo é um híbrido entre autômato celular estocástico e modelo baseado em agentes: o espaço é discretizado em células, enquanto cada filhote mantém estado, parâmetros e trajetória próprios.")

    doc.add_paragraph("3.1 Entidades, escalas e condições de contorno", style="Heading 2")
    add_text(doc, "O eixo x é transversal à praia e cresce da duna para o oceano. O eixo y acompanha a costa. A configuração de demonstração utiliza células de 0,5 m e passos de 5 s. A linha d’água é uma função da posição longitudinal, e sua ultrapassagem encerra a trajetória com o desfecho ‘mar’. As bordas laterais são refletoras no exemplo, evitando o reaparecimento toroidal incompatível com fontes localizadas. A borda terrestre é um desfecho específico, e indivíduos sem desfecho até o tempo máximo são mantidos como censurados.")
    add_table(
        doc,
        ["Símbolo/arquivo", "Significado", "Unidade/escala"],
        [
            ["Δx", "Tamanho da célula", "m"],
            ["Δt", "Duração do passo", "s"],
            ["x, y", "Posição transversal e longitudinal", "m"],
            ["waterline_x(y)", "Linha d’água local", "m"],
            ["Qᵢ", "Reserva energética individual", "adimensional ou calibrada"],
            ["hᵢ,t", "Direção anterior", "vetor unitário"],
            ["τᵢ", "Temperatura da Softmax", "adimensional"],
        ],
        [1.5, 3.1, 1.75],
    )

    doc.add_paragraph("3.2 Conversão e decomposição do campo luminoso", style="Heading 2")
    add_text(doc, "A magnitude direcional m é convertida em radiância relativa. A transformação preserva razões de brilho, mas não converte o SQM em lux nem em uma resposta espectral específica da tartaruga.")
    add_formula(doc, [("R", None), ("d", "sub"), (" = exp{−0,4 ln(10) [m", None), ("d", "sub"), (" − m", None), ("ref", "sub"), ("]}", None)])
    add_text(doc, "A pista marítima é normalizada pela mediana das leituras na direção do oceano:")
    add_formula(doc, [("B", None), ("sea", "sub"), ("(y) = R", None), ("ocean", "sub"), ("(y) / mediana[R", None), ("ocean", "sub"), ("]", None)])
    add_text(doc, "A componente artificial é obtida operacionalmente pelo excesso de brilho nas direções terrestre e laterais em relação à leitura do oceano:")
    add_formula(doc, [("A", None), ("x", "sub"), ("(y) = −[R", None), ("dune", "sub"), ("(y) − R", None), ("ocean", "sub"), ("(y)]", None), ("+", "sub"), (" / R̃", None), ("ocean", "sub")])
    add_formula(doc, [("A", None), ("y", "sub"), ("(y) = {[R", None), ("north", "sub"), (" − R", None), ("ocean", "sub"), ("]", None), ("+", "sub"), (" − [R", None), ("south", "sub"), (" − R", None), ("ocean", "sub"), ("]", None), ("+", "sub"), ("} / R̃", None), ("ocean", "sub")])
    add_text(doc, "O vetor é atenuado transversalmente e pela oclusão aparente da duna. Os parâmetros de cenário representam redução de intensidade, blindagem e um peso espectral proxy. O peso âmbar não é uma constante biológica universal; deve ser substituído por calibração espectral da espécie e das luminárias locais.")

    doc.add_paragraph("3.3 Regra de movimento", style="Heading 2")
    add_text(doc, "Para cada célula da vizinhança de Moore, incluindo a possibilidade de permanecer parada, calcula-se uma utilidade. As pistas marítima, artificial, de silhueta e de declive são vetores, e sua contribuição depende do alinhamento com a direção candidata.")
    add_formula(doc, [("U", None), ("i,k", "sub"), (" = κ", None), ("sea,i", "sub"), (" m", None), ("k", "sub"), ("·S + κ", None), ("art,i", "sub"), (" m", None), ("k", "sub"), ("·A + κ", None), ("dune,i", "sub"), (" m", None), ("k", "sub"), ("·D + κ", None), ("slope,i", "sub"), (" m", None), ("k", "sub"), ("·G + ρ", None), ("i", "sub"), (" m", None), ("k", "sub"), ("·h", None), ("i,t−1", "sub"), (" − c", None), ("stay", "sub"), (" I", None), ("k", "sub")])
    add_text(doc, "A probabilidade de selecionar o movimento k é:")
    add_formula(doc, [("P", None), ("i,k", "sub"), (" = exp(U", None), ("i,k", "sub"), ("/τ", None), ("i", "sub"), (") / Σ", None), ("j∈Nᵢ", "sub"), (" exp(U", None), ("i,j", "sub"), ("/τ", None), ("i", "sub"), (")", None)])
    add_text(doc, "A temperatura τ controla a aleatoriedade: valores altos aproximam as escolhas entre si; valores baixos concentram a probabilidade nos movimentos de maior utilidade. A persistência ρ é separada da sensibilidade luminosa, evitando que um único parâmetro represente processos diferentes.")

    doc.add_paragraph("3.4 Heterogeneidade, energia e predação", style="Heading 2")
    add_text(doc, "Os coeficientes individuais positivos são amostrados por distribuições lognormais, e a persistência por uma normal truncada. A reserva energética é reduzida por tempo, distância e declividade:")
    add_formula(doc, [("Q", None), ("i", "sub"), ("(t+Δt) = Q", None), ("i", "sub"), ("(t) − c", None), ("b", "sub"), ("Δt − c", None), ("m", "sub"), ("d", None), ("i", "sub"), ("[1 + c", None), ("s", "sub"), ("|sen(s(y))|]", None)])
    add_text(doc, "A predação é modelada como risco por tempo de exposição, em vez de exigir colisão exata com um predador em uma célula arbitrária:")
    add_formula(doc, [("p", None), ("pred", "sub"), (" = 1 − exp[−λ(x,y,t) Δt]", None)])
    add_text(doc, "Sem observações locais, os parâmetros de energia e predação devem permanecer desativados ou ser explorados em análise de sensibilidade. Nesta demonstração, eles são cenários e não estimativas de mortalidade.")

    doc.add_paragraph("4. Calibração, verificação e validação", style="Heading 1")
    doc.add_paragraph("4.1 Verificação computacional", style="Heading 2")
    add_text(doc, "O projeto inclui testes de conservação de indivíduos, reprodutibilidade com sementes fixas, sinais opostos das pistas marítima e artificial e conferência do número de medições reais. A estabilidade numérica da Softmax é obtida pela subtração do maior logit antes da exponenciação.")

    doc.add_paragraph("4.2 Calibração ambiental", style="Heading 2")
    add_text(doc, "Os nove pontos são projetados em um eixo local ao longo da costa. A validação cruzada exclui um local por vez e interpola seu brilho a partir dos demais. O erro quadrático médio global foi de " + fmt_num(light_validation["overall_rmse_mag_arcsec2"], 3) + " mag/arcsec². O horizonte do oceano apresentou RMSE de " + fmt_num(light_validation["by_direction"]["ocean"]["rmse_mag_arcsec2"], 3) + ", enquanto a direção da duna apresentou " + fmt_num(light_validation["by_direction"]["dune"]["rmse_mag_arcsec2"], 3) + ". Esses valores avaliam a interpolação longitudinal, não a decomposição física entre luz natural e artificial.")

    doc.add_paragraph("4.3 Calibração comportamental", style="Heading 2")
    add_text(doc, "A calibração minimiza uma soma ponderada de erros entre resumos observados e simulados: proporção que alcança o mar, desorientação operacional, tempo mediano, eficiência mediana e desvio angular inicial. Utiliza-se evolução diferencial e sementes comuns entre avaliações. Os parâmetros ajustados na demonstração rápida foram:")
    add_table(
        doc,
        ["Parâmetro", "Estimativa demonstrativa", "Interpretação"],
        [
            ["κsea", fmt_num(calibration["parameters"]["kappa_sea"], 3), "Resposta à pista marítima"],
            ["κartificial", fmt_num(calibration["parameters"]["kappa_artificial"], 3), "Resposta ao campo artificial"],
            ["ρ", fmt_num(calibration["parameters"]["persistence"], 3), "Persistência do rumo"],
            ["τ", fmt_num(calibration["parameters"]["temperature"], 3), "Aleatoriedade da Softmax"],
        ],
        [1.45, 1.65, 3.45],
    )
    add_note_box(doc, "Convergência", "O modo rápido encerrou pelo limite reduzido de iterações, não pelo critério completo de convergência. As métricas ajustadas ficaram próximas às sintéticas, mas os parâmetros não devem ser interpretados biologicamente. A análise final deve ser executada sem --quick, com mais indivíduos, repetições e diagnóstico de identificabilidade.")
    add_figure(doc, figs / "calibration_fit.png", "Figura 2. Comparação entre métricas sintéticas usadas na calibração e métricas produzidas pelo ajuste rápido. O gráfico demonstra o fluxo computacional; não representa validação de campo.", 5.9)

    doc.add_paragraph("4.4 Validação comportamental independente", style="Heading 2")
    obs = behavior_validation["observed_metrics"]
    sim = behavior_validation["simulated_metrics"]
    add_text(doc, "Um segundo conjunto sintético, gerado com sementes independentes, foi reservado para validação. A chegada ao mar foi " + fmt_pct(obs["sea_arrival_rate"]) + " no conjunto de validação e " + fmt_pct(sim["sea_arrival_rate"]) + " na simulação calibrada. A desorientação foi " + fmt_pct(obs["disorientation_rate"]) + " e " + fmt_pct(sim["disorientation_rate"]) + ", respectivamente. O tempo mediano até o mar diferiu em " + fmt_num(sim["median_time_to_sea_s"] - obs["median_time_to_sea_s"], 0) + " s. Essa etapa verifica a separação lógica entre calibração e validação, mas continua sintética.")
    add_table(
        doc,
        ["Métrica", "Validação sintética", "Modelo calibrado", "Diferença"],
        [
            ["Chegada ao mar", fmt_pct(obs["sea_arrival_rate"]), fmt_pct(sim["sea_arrival_rate"]), fmt_num(sim["sea_arrival_rate"]-obs["sea_arrival_rate"], 3)],
            ["Desorientação", fmt_pct(obs["disorientation_rate"]), fmt_pct(sim["disorientation_rate"]), fmt_num(sim["disorientation_rate"]-obs["disorientation_rate"], 3)],
            ["Tempo mediano (s)", fmt_num(obs["median_time_to_sea_s"], 0), fmt_num(sim["median_time_to_sea_s"], 0), fmt_num(sim["median_time_to_sea_s"]-obs["median_time_to_sea_s"], 0)],
            ["Eficiência mediana", fmt_num(obs["median_efficiency"], 3), fmt_num(sim["median_efficiency"], 3), fmt_num(sim["median_efficiency"]-obs["median_efficiency"], 3)],
            ["Erro angular médio (°)", fmt_num(obs["mean_abs_initial_heading_deg"], 1), fmt_num(sim["mean_abs_initial_heading_deg"], 1), fmt_num(sim["mean_abs_initial_heading_deg"]-obs["mean_abs_initial_heading_deg"], 1)],
        ],
        [1.8, 1.55, 1.55, 1.25],
    )

    doc.add_paragraph("5. Experimentos de Monte Carlo", style="Heading 1")
    add_text(doc, "Seis cenários foram avaliados. Cada cenário da execução de referência utilizou oito repetições de 30 indivíduos. Os intervalos nas figuras representam a incerteza da média entre repetições. Para um estudo final, recomenda-se aumentar as repetições até que médias e intervalos se estabilizem.")
    add_table(
        doc,
        ["Cenário", "Componente artificial", "Outras alterações"],
        [
            ["Referência", "100%", "Pista marítima padrão"],
            ["Luzes desligadas", "0%", "Sem componente artificial"],
            ["Blindagem 70%", "30% visível", "Blindagem direcional"],
            ["Âmbar (proxy)", "Peso 45%", "Proxy espectral"],
            ["Lua mais intensa", "100%", "Pista marítima × 1,35"],
            ["Mitigação combinada", "Escala 55%, peso 45%, blindagem 70%", "Pista marítima × 1,15 e duna +0,6 m"],
        ],
        [1.8, 2.1, 2.55],
    )

    doc.add_paragraph("6. Resultados demonstrativos", style="Heading 1")
    rows = []
    for r in scenario.itertuples(index=False):
        rows.append([
            SCENARIO_LABELS.get(r.scenario, r.scenario),
            fmt_pct(r.sea_arrival_rate_mean),
            fmt_pct(r.disorientation_rate_mean),
            fmt_pct(r.predation_rate_mean),
            fmt_pct(r.censored_rate_mean),
            fmt_num(r.median_time_to_sea_s_mean, 0),
            fmt_num(r.median_efficiency_mean, 3),
        ])
    add_table(doc, ["Cenário", "Mar", "Desorient.", "Predação", "Censura", "Tempo (s)", "Eficiência"], rows, [1.65, 0.72, 0.85, 0.78, 0.72, 0.78, 0.83], font_size=8.6)
    add_text(doc, "Na condição de referência, a chegada ao mar foi " + fmt_pct(float(scenario.loc[scenario.scenario=="baseline", "sea_arrival_rate_mean"].iloc[0])) + ", com desorientação operacional de " + fmt_pct(float(scenario.loc[scenario.scenario=="baseline", "disorientation_rate_mean"].iloc[0])) + ". A remoção da componente artificial elevou a chegada ao mar para " + fmt_pct(float(scenario.loc[scenario.scenario=="lights_off", "sea_arrival_rate_mean"].iloc[0])) + " e reduziu o tempo mediano de aproximadamente 700 s para 426 s. A mitigação combinada produziu o maior desempenho médio, mas sua semelhança com o cenário de luzes desligadas não deve ser convertida em recomendação quantitativa sem dados de campo.")
    add_figure(doc, figs / "scenario_sea_arrival.png", "Figura 3. Proporção média que alcançou a linha d’água em cada cenário demonstrativo. Barras de erro: intervalo aproximado de 95% da média entre repetições.", 6.25)
    add_figure(doc, figs / "scenario_disorientation.png", "Figura 4. Desorientação operacional por cenário. A definição combina erro angular, eficiência de trajetória e saídas não marítimas; deve ser alinhada ao protocolo de campo antes da calibração real.", 6.25)
    add_figure(doc, figs / "sample_tracks_baseline.png", "Figura 5. Amostra de trajetórias no cenário de referência. O oceano está à direita. O perfil transversal e os ninhos desta figura são demonstrativos.", 6.25)

    doc.add_paragraph("6.1 Sensibilidade global", style="Heading 2")
    sea_sens = sensitivity[sensitivity.output == "sea_arrival_rate"].set_index("parameter")["spearman_rho"].to_dict()
    dis_sens = sensitivity[sensitivity.output == "disorientation_rate"].set_index("parameter")["spearman_rho"].to_dict()
    add_text(doc, "Na amostra exploratória por hipercubo latino, κsea apresentou correlação de Spearman de " + fmt_num(sea_sens["kappa_sea"], 2) + " com a chegada ao mar, enquanto κartificial apresentou " + fmt_num(sea_sens["kappa_artificial"], 2) + ". Para a desorientação, os sinais foram opostos: " + fmt_num(dis_sens["kappa_sea"], 2) + " e " + fmt_num(dis_sens["kappa_artificial"], 2) + ". O pequeno número de amostras da demonstração não substitui uma análise de Sobol ou uma amostragem ampliada.")
    add_figure(doc, project / "outputs_demo" / "sensitivity" / "sensitivity_sea_arrival.png", "Figura 6. Correlações de Spearman entre parâmetros calibráveis e chegada ao mar na análise de sensibilidade demonstrativa.", 5.7)

    doc.add_paragraph("7. Discussão", style="Heading 1")
    add_text(doc, "A principal melhoria conceitual é a representação direcional da luz. No modelo reformulado, a pista marítima aponta para o oceano e a componente artificial derivada do contraste terrestre pode apontar para a duna ou ao longo da costa. A poluição luminosa deixa de ser um número sem unidade que multiplica um gradiente escolhido arbitrariamente e passa a ser vinculada a um levantamento espacial reproduzível.")
    add_text(doc, "A segunda melhoria é a separação de mecanismos. Sensibilidade às pistas, persistência e aleatoriedade possuem parâmetros distintos. Isso reduz a confusão estrutural do modelo anterior, embora não elimine problemas de identificabilidade: combinações diferentes podem produzir resumos semelhantes. Por esse motivo, trajetórias completas, e não apenas proporções finais, são necessárias para estimar parâmetros com maior precisão.")
    add_text(doc, "A terceira melhoria é a distinção entre chegada ao mar e sobrevivência. O desfecho principal representa a conclusão da travessia. Ele não estima a sobrevivência até a fase adulta nem explica a razão demográfica frequentemente citada de um adulto para muitos nascimentos. Predação, exaustão e censura são registradas separadamente.")
    add_text(doc, "Os resultados demonstrativos são coerentes com a estrutura imposta: atenuar o vetor artificial favorece a pista marítima. Contudo, essa coerência interna não constitui validação ecológica. O efeito quantitativo depende da largura real da praia, da altura da duna, da localização dos ninhos, do espectro luminoso, das condições lunares, da velocidade e das trajetórias observadas.")

    doc.add_paragraph("7.1 Limitações", style="Heading 2")
    add_bullet(doc, "O levantamento SQM é de uma noite e fornece brilho de banda larga, não radiância espectral ponderada pela visão da espécie.")
    add_bullet(doc, "A separação entre componente natural e artificial usa o horizonte oceânico como referência operacional.")
    add_bullet(doc, "O perfil, os ninhos e os obstáculos do exemplo não são dados topográficos de Delray Beach.")
    add_bullet(doc, "A calibração e a validação comportamental apresentadas são sintéticas.")
    add_bullet(doc, "O módulo de predação é um risco agregado e requer dados próprios para interpretação biológica.")
    add_bullet(doc, "As categorias ‘âmbar’ e ‘lua mais intensa’ são proxies de cenário, não conversões universais.")

    doc.add_paragraph("8. Protocolo para calibração integral", style="Heading 1")
    add_text(doc, "Para transformar esta versão em modelo calibrado para uma praia específica, o estudo deve executar as seguintes etapas:")
    for item in [
        "levantar perfis transversais, linha d’água, duna e declividade no período das emergências;",
        "medir o campo luminoso em direções e alturas documentadas, incluindo espectro, Lua, nuvens e maré;",
        "georreferenciar ninhos e digitalizar trajetórias com vídeo infravermelho;",
        "separar noites ou ninhos entre conjuntos de calibração e validação;",
        "executar a calibração completa, verificar convergência e identificabilidade;",
        "validar em dados independentes e testar a transferência para outros setores/noites;",
        "realizar Monte Carlo e análise global de sensibilidade antes de comparar intervenções.",
    ]:
        add_bullet(doc, item)
    add_text(doc, "O pacote computacional acompanha modelos de CSV, protocolo de campo, testes, comandos de calibração, validação, cenários e figuras. Todas as sementes e suposições são armazenadas em YAML.")

    doc.add_paragraph("9. Conclusões", style="Heading 1")
    add_text(doc, "A reformulação corrige o sinal do campo luminoso, atribui escalas físicas ao espaço e ao tempo, separa processos comportamentais, mantém indivíduos censurados no denominador e introduz calibração e validação explícitas. O estudo de caso utiliza medições reais de brilho direcional para construir o ambiente, mas não reivindica calibração biológica integral enquanto geometria e trajetórias locais não forem incorporadas. A contribuição principal desta versão é um arcabouço reproduzível no qual hipóteses sobre iluminação, blindagem, espectro e condições naturais podem ser testadas sem confundir parâmetros abstratos com intensidades reais.")

    doc.add_paragraph("Disponibilidade do código e reprodução", style="Heading 1")
    add_text(doc, "O pacote acompanha os módulos do modelo, configuração de Delray Beach, dados de entrada, testes, resultados demonstrativos e um guia de execução. O fluxo rápido é iniciado por scripts/run_delray_demo.sh ou scripts\\run_delray_demo.bat. Para a análise final, a calibração deve ser executada sem a opção --quick e com dados de campo separados para validação.")

    doc.add_paragraph("Referências Bibliográficas", style="Heading 1")
    refs = [
        "Celano, L., Sullivan, C., Field, A., & Salmon, M. (2018). Seafinding revisited: how hatchling marine turtles respond to natural lighting at a nesting beach. Journal of Comparative Physiology A, 204, 1007–1015. DOI: 10.1007/s00359-018-1299-4.",
        "Grimm, V., Railsback, S. F., Vincenot, C. E., et al. (2020). The ODD Protocol for Describing Agent-Based and Other Simulation Models: A Second Update to Improve Clarity, Replication, and Structural Realism. Journal of Artificial Societies and Social Simulation, 23(2), 7. DOI: 10.18564/jasss.4259.",
        "Hirama, S., Sylvia, A., Long, T., Trindell, R., & Witherington, B. (2022). Light brightness data near sea turtle nests as measured from the horizon and zenith using a Sky Quality Meter. Data in Brief, 43, 108430. DOI: 10.1016/j.dib.2022.108430.",
        "Lohmann, K. J., Witherington, B. E., Lohmann, C. M. F., & Salmon, M. (1997). Orientation, navigation, and natal beach homing in sea turtles. In The Biology of Sea Turtles (Vol. 1). CRC Press.",
        "Salmon, M. (2003). Artificial night lighting and sea turtles. Biologist, 50(4), 163–168.",
        "Salmon, M., Wyneken, J., Fritz, E., & Lucas, M. (1992). Seafinding by hatchling sea turtles: role of brightness, silhouette and beach slope as orientation cues. Behaviour, 122, 56–77.",
        "Schiff, J. L. (2008). Cellular Automata: A Discrete View of the World. Wiley-Interscience.",
        "Stanley, T. R., White, J. M., Teel, S., & Nicholas, M. (2020). Brightness of the Night Sky Affects Loggerhead (Caretta caretta) Sea Turtle Hatchling Misorientation but Not Nest Site Selection. Frontiers in Marine Science, 7, 221. DOI: 10.3389/fmars.2020.00221.",
        "Witherington, B. E., & Martin, R. E. (2000). Understanding, Assessing, and Resolving Light-Pollution Problems on Sea Turtle Nesting Beaches. Florida Marine Research Institute.",
    ]
    for ref in refs:
        add_reference(doc, ref)

    doc.add_paragraph("Apêndice A — Comandos essenciais", style="Heading 1")
    commands = [
        ("Validar entradas", "python -m turtle_beach_model.cli validate-data --config configs/minha_praia.yaml"),
        ("Resumir trajetórias", "python -m turtle_beach_model.cli summarize-observed --config configs/minha_praia.yaml --coordinates data/coordenadas.csv --output data/resumos.csv"),
        ("Calibrar", "python -m turtle_beach_model.cli calibrate --config configs/minha_praia.yaml --observed data/calibracao.csv --output outputs/calibration.json"),
        ("Validar", "python -m turtle_beach_model.cli validate-calibration --config configs/minha_praia.yaml --observed data/validacao.csv --calibration outputs/calibration.json --output outputs/validation.json --replicates 30"),
        ("Executar cenários", "python -m turtle_beach_model.cli experiment --config configs/minha_praia.yaml --calibration outputs/calibration.json --output-dir outputs/experiments --replicates 100 --n-turtles 50"),
        ("Gerar figuras", "python -m turtle_beach_model.cli plot --config configs/minha_praia.yaml --experiment-dir outputs/experiments --calibration outputs/calibration.json"),
        ("Sensibilidade", "python -m turtle_beach_model.cli sensitivity --config configs/minha_praia.yaml --output-dir outputs/sensitivity --samples 128 --n-turtles 50"),
    ]
    for label, command in commands:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(2)
        lead = p.add_run(label + ": ")
        lead.bold = True
        lead.font.name = "Arial"
        lead.font.size = Pt(9.2)
        r = p.add_run(command)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(7.8)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--project-root", default=str(Path(__file__).resolve().parents[1]))
    parser.add_argument("--output", required=True)
    args = parser.parse_args()
    generate(Path(args.project_root).resolve(), Path(args.output).resolve())


if __name__ == "__main__":
    main()
